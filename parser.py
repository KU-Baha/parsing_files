#!/usr/bin/python3
# -*- coding: utf-8 -*-

import sys as Sys
import os as Os
import pandas as Ps
import pathlib as Pt
import json as Jn
import PyPDF2 as Pp2
import camelot as Ct
import openpyxl as Ox
import zipfile as Zf
import re as Re
import docx as Dc
import io as Io
import subprocess as Sp
import numpy as Np
from xml.dom import minidom
import multiprocessing as Mp
import math as Mt
import time as Tm

Sys.setrecursionlimit(15000)


class Parser:
    debug = False
    params = None
    manager = None
    managerProcess = None
    cutPartsPdf = Mt.ceil(int(Mp.cpu_count()) / 1)
    minimumColumns = 3

    pathHandlerSite = '/opt/php81/bin/php /var/www/www-root/data/www/costana.testdom.online/artisan parser_check:check '
    paramsHandlerSite = ["fileId", "objectId", "projectId", "status", "statusText"]
    directory = {"in": None, "out": None}
    file = {"in": None, "out": None}
    extentions = [".xls", ".doc", ".xlsx", ".docx", ".pdf"]
    sizesPages = {
        'a4': {'width': 595, 'height': 842},
        'a3': {'width': 842, 'height': 1191}
    }
    data = []
    dataOutput = []
    dataColumnsSearch = []
    dataColumnsSearchLength = 0
    dataColumns = {
        'position': ['поз', 'позиция', '#', '№'],
        'name': ['наим', 'наименование', 'наз', 'название'],
        'code': ['код', 'артикул', 'арт'],
        'type': ['тип', 'марка'],
        'customer': ['завод', 'изготовитель', 'поставщик', 'фирма'],
        'unit': ['ед', 'единица', 'имерения'],
        'count': ['кол', 'кол-во', 'количество'],
        'weight': ['масса', 'вес'],
        'note': ['примечание']
    }
    dataColumnsResult = [
        'name',
        'code',
        'type',
        'customer',
        'unit',
        'count',
        'price_work',
        'full_price_work',
        'price_material',
        'full_price_material',
        'note'
    ]

    def __init__(self, directory, file, params):
        print(params)
        errors = []

        if params is not None:
            self.params = params

        if directory is not None:
            if directory["in"] is not None:
                self.directory["in"] = directory["in"]
            else:
                errors.append("ERR_DIR_IN")

            if directory["out"] is not None:
                self.directory["out"] = directory["out"]
            else:
                errors.append("ERR_DIR_OUT")
        else:
            errors.append("ERR_DIR")

        if file is not None:
            if file["in"] is not None:
                self.file["in"] = file["in"]
            else:
                errors.append("ERR_FILE_IN")

            if directory["out"] is not None:
                self.file["out"] = file["out"]
            else:
                errors.append("ERR_FILE_OUT")
        else:
            errors.append("ERR_FILE")

        if len(errors):
            self.errorsSignal(errors)
        else:
            self.checkDirectory()

    def checkDirectory(self):
        errors = []

        if Os.path.exists(self.directory["in"]) == False:
            errors.append("ERR_DIR_IN_EXIST" + self.directory["in"])

        if Os.path.exists(self.directory["out"]) == False:
            Os.mkdir(self.directory["out"])

        if len(errors):
            self.errorsSignal(errors)
        else:
            self.checkFile()

    def checkFile(self):

        errors = []

        if Os.path.exists(self.file["in"]) == False:
            errors.append("ERR_FILE_IN_EXIST " + self.file["in"])

        if Os.path.exists(self.file["out"]) == False:
            open(self.file["out"], "w").close()

        if len(errors):
            self.errorsSignal(errors)
        else:
            self.checkFileExtension()

    def checkFileExtension(self):

        errors = []
        extention = Pt.Path(self.file["in"]).suffix

        if extention in self.extentions == False:
            errors.append("ERR_FILE_EXTENTION " + self.file["in"])

        if len(errors):
            self.errorsSignal(errors)
        else:
            self.handlerFile(extention)

    def handlerFile(self, extention):

        if extention == '.xls' or extention == '.xlsx':
            self.ExcelHandler(extention)
        elif extention == '.doc' or extention == '.docx':
            self.WordHandler(extention)
        elif extention == '.pdf':
            self.PdfHandler()

    def PsClean(self, df):

        df.replace(to_replace='\(cid\:[0-9]+\)', value='', inplace=True, regex=True)
        df.replace(to_replace='[\\n\\r\\t]', value='', inplace=True, regex=True)
        df.replace(to_replace='None', inplace=True, value=Np.nan)
        df.replace(Np.nan, '', regex=True, inplace=True)
        df.dropna(axis=0, inplace=True)
        df.dropna(axis=1, inplace=True)
        df.drop_duplicates(keep="first", inplace=True)

        return df

    def cutPdfPages(self, list, num):

        part_len = Mt.ceil(len(list) / num)

        return [list[part_len * k:part_len * (k + 1)] for k in range(0, num)]

    def PdfHandlerFilterPageSize(self, width, height, rotate, numPage):

        success = False
        width = round(float(width))
        height = round(float(height))

        if height < width:
            width, height = height, width

        for index in self.sizesPages:

            sizePage = self.sizesPages[index]

            if 'width' in sizePage and 'height' in sizePage:

                if width <= sizePage['width'] and height <= sizePage['height']:
                    success = True
                    break

        return success

    def PdfHandler(self):

        fileIn = Pp2.PdfFileReader(self.file["in"], strict=False)

        if fileIn is not None and fileIn.getNumPages() > 0:

            pdfPages = []
            pdf = Pp2.PdfFileReader(self.file["in"], strict=False)
            numPages = pdf.getNumPages();

            for numPage in range(0, numPages):

                pdfPage = pdf.getPage(numPage)

                if self.PdfHandlerFilterPageSize(pdfPage.mediaBox.getWidth(), pdfPage.mediaBox.getHeight(),
                                                 pdfPage.get('/Rotate'), numPage):
                    pdfPages.append(numPage)

            numPages = len(pdfPages)

            if numPages:

                print('PdfHandler -> ', numPages)

                pdfPages = self.cutPdfPages(pdfPages, self.cutPartsPdf)

                if len(pdfPages):
                    with Mp.Manager() as manager:

                        self.manager = manager.list()
                        self.managerProcess = list()

                        for processPage in pdfPages:
                            if len(processPage):
                                self.PdfThreadPagesHandlerStart(processPage, numPages)

                        if len(self.managerProcess):
                            for process in self.managerProcess:
                                process.join()
            else:
                self.errorsSignal('ERR_FILE_ERROR')

        else:
            self.errorsSignal('ERR_FILE_ERROR')

    def PdfThreadPagesHandlerStart(self, threadPage, numPages):
        th = Mp.Process(target=self.PdfThreadPagesHandler, args=(threadPage, numPages))
        th.start()
        self.managerProcess.append(th)

    def PdfThreadPagesHandler(self, pages, numPages):
        for page in pages:
            print('PdfThreadPagesHandler -> ', page)
            self.PdfPageHandler(self.file["in"], page)
            self.PdfThreadPagesHandlerFinish(numPages)

    def PdfThreadPagesHandlerFinish(self, numPages):

        if len(self.manager) == numPages:

            result = list()

            for item in self.manager:
                if item is not None:
                    result.append(item)

            self.data = result

            self.handlerData()

    def PdfPageHandler(self, path, page):

        tables = Ct.read_pdf(path, flavor='stream', edge_tol=500, row_tol=15, pages=str(page))

        if tables:
            result = None

            for table in tables:

                data = table.df
                data = self.PsClean(data)
                data = Jn.loads(data.to_json(orient="values"))

                if data:

                    data = self.prepareData(data)

                    if data:
                        if result is not None:
                            result += data
                        else:
                            result = data

                        # Ct.plot(table, kind='textedge').show()

            self.manager.insert(page, result)

    def WordHandler(self, extention):

        if (extention == '.doc'):
            Sp.call(
                ['soffice', '--headless', '--convert-to', 'docx', self.file["in"], '--outdir', self.directory["in"]])
            old_path = self.file["in"]
            path = Pt.Path(self.file["in"])
            self.file["in"] = path.with_name(path.stem + ".docx")

            # Os.unlink(old_path)

        doc = Dc.Document(self.file["in"])
        tables = []

        for section in doc.sections:

            if section.header is not None:
                if section.header.tables is not None:
                    for header in section.header.tables:
                        tables.append(header)

        if len(doc.tables):
            for table in doc.tables:
                tables.append(table)

        # for section in doc.sections:
        #    if section.footer is not None:
        #        if section.footer.tables is not None:
        #            for footer in section.footer.tables:
        #                tables.append(footer)

        if len(tables):
            self.WordLoadTables(tables)

    def WordLoadTables(self, tables):

        data = {}
        irow = 0

        # print('WordLoadTables')

        if len(tables):
            for table in tables:
                if len(table.rows):
                    for row in table.rows:

                        # if irow < 10:

                        if len(row.cells):

                            vcell = {}
                            icell = 0

                            for cell in row.cells:
                                vcell[icell] = cell.text
                                icell += 1

                            if len(vcell):
                                if irow == 0:
                                    data[irow] = ";".join('"' + str(x) + '"' for x in range(0, len(vcell)))
                                    irow += 1

                                data[irow] = ";".join('"' + str(x) + '"' for x in vcell.values())
                                irow += 1

        if len(data):

            data = "\r\n".join(str(x) for x in data.values())
            data = Ps.read_csv(Io.StringIO(data), sep=';')
            data = self.PsClean(data)
            data = Jn.loads(data.to_json(orient="values"))

            if data:
                data = self.prepareData(data)

                if data:
                    self.data.append(data)

        self.handlerData()

    def ExcelHandler(self, extention):

        sheets_names = self.ExcelGetSheets(self.file["in"])

        if len(sheets_names):
            for sheet_name in sheets_names:

                data = self.ExcelLoadSheet(sheet_name)
                data = self.PsClean(data)
                data = Jn.loads(data.to_json(orient="values"))

                if data:
                    data = self.prepareData(data)

                    if data:
                        self.data.append(data)

        self.handlerData()

    def ExcelLoadSheet(self, name):

        sheet = Ps.read_excel(self.file["in"], engine="openpyxl", sheet_name=str(name))

        return sheet

    def ExcelGetSheets(self, path):

        sheets = []

        with Zf.ZipFile(path, 'r') as zip_ref:
            xml = zip_ref.read("xl/workbook.xml").decode("utf-8")

            for s_tag in Re.findall("<sheet [^>]*", xml):
                sheets.append(Re.search('name="[^"]*', s_tag).group(0)[6:])

        return sheets

    def prepareData(self, data):

        current_data = data

        for indexRow, row in enumerate(current_data):
            if len(row) < self.minimumColumns:
                current_data.remove(row)
                self.prepareData(data)
                break

            else:
                if not self.prepareRowAfter(row):
                    current_data.remove(row)
                    self.prepareData(data)
                    break

        return current_data;

    def prepareRowBefore(self, row):

        var = 0

        for indexCol, col in enumerate(row):
            if col is not None:
                var += 1

        return var

    def prepareRowAfter(self, row):

        var = 0

        for indexCol, col in enumerate(row):
            tmp = Re.sub(r"[^А-Яа-яёЁ]", "", str(col), 0, Re.MULTILINE)  # [^A-Za-zА-Яа-я]

            if len(tmp):
                var += 1

        return var

    def preparePageTableHead(self, index, page):

        mergeRow = False
        lastIndexRow = None
        lastRow = None

        for indexRow, row in enumerate(page):
            if mergeRow == True:
                newRow = self.mergeRow(lastRow, row)
                page.insert(lastIndexRow, newRow)
                del page[lastIndexRow + 1]
                del page[lastIndexRow + 1]
                self.preparePageTableHead(index, page)
                break

            for indexCell, cell in enumerate(row):
                cell = str(cell).strip()

                if len(cell):
                    symbol = cell[-1]

                    if symbol == ',' or symbol == '-':

                        entry = False
                        columns = self.dataColumns.copy()

                        for column in dict(columns):
                            for columnElement in columns[column]:
                                result = Re.match(r"^" + columnElement.lower(), str(cell).lower())

                                if result is not None:
                                    entry = True
                                    break

                        if entry == True:
                            mergeRow = True
                            break

            lastIndexRow = indexRow
            lastRow = row

    def preparePageTableBody(self, index, page):

        mergeRow = False
        lastIndexRow = None
        lastRow = None

        for indexRow, row in enumerate(page):
            if mergeRow == True:
                newRow = self.mergeRow(lastRow, row)
                page.insert(lastIndexRow, newRow)
                del page[lastIndexRow + 1]
                del page[lastIndexRow + 1]
                self.preparePageTableHead(index, page)
                break

            for indexCell, cell in enumerate(row):
                cell = cell.strip()

                if len(cell):
                    symbol = cell[-1]

                    if symbol == ',' or symbol == '-':
                        mergeRow = True
                        break

            lastIndexRow = indexRow
            lastRow = row

    def mergeRow(self, firstRow, lastRow):

        row = []

        if firstRow is not None and lastRow is not None:
            for index, value in enumerate(firstRow):
                if lastRow[index] is not None:
                    if len(value):
                        symbol = value[-1]

                        if symbol == ',' or symbol == '-':
                            row.append(value + lastRow[index])
                        else:
                            row.append(value + " " + lastRow[index])
                    else:
                        row.append(lastRow[index])
                else:
                    row.append(value)

        return row

    def searchHead(self, row, step=1):

        output = {}
        hiddenColumns = []
        columns = self.dataColumns.copy()

        for indexCell, valueCell in enumerate(row):
            valueCell = str(valueCell)

            if len(valueCell):
                for column in dict(columns):
                    for columnElement in columns[column]:

                        valueCell = valueCell.strip()
                        valueCell = Re.sub(r"[^\w]|[_][^,]", "", valueCell, 0, Re.MULTILINE)

                        if valueCell is not None:
                            result = Re.match(r"^" + columnElement.lower(), valueCell.lower())

                            if result is not None:
                                output[column] = indexCell

                                if column in columns:
                                    del columns[column]

        if len(output):
            return {"output": output, "length": len(row)}

    def searchHeadAfter(self, result):

        if result is not None and "output" in result and "length" in result:
            if result["output"] is not None and result["length"] is not None:
                if result["length"] >= self.minimumColumns and len(result["output"]) >= self.minimumColumns:
                    self.dataColumnsSearch = result["output"]
                    self.dataColumnsSearchLength = result["length"]

    def searchHeadReset(self):

        self.dataColumnsSearch = []
        self.dataColumnsSearchLength = 0

    def searchData(self, row):

        result = {}

        if (len(row) == self.dataColumnsSearchLength):
            for column in self.dataColumnsSearch:
                index = self.dataColumnsSearch[column]

                if row[index] is not None:
                    result[column] = row[index]

        if len(result):
            return result

    def prepareSignal(self, status, statusText=[""]):

        command = ""

        if self.pathHandlerSite is not None and len(
                self.pathHandlerSite) and self.paramsHandlerSite is not None and len(self.paramsHandlerSite):
            command = self.pathHandlerSite

            for param in self.paramsHandlerSite:

                value = ""

                if param in self.params:
                    value = self.params[param]

                if param == "status":
                    value = status

                if param == "statusText":
                    value = "".join(statusText)

                command += "=".join(["--" + param, '"' + str(value) + '"'])
                command += " "

        # print(command)

        return command

    def successSignal(self):

        print('status -> success')

        if self.debug is not None and self.debug == False:
            Os.system(self.prepareSignal(1))

        Sys.exit()

    def errorsSignal(self, errors):

        print('status -> errors', errors)

        if self.debug is not None and self.debug == False:
            Os.system(self.prepareSignal(0, errors))

        Sys.exit()

    def handlerData(self):

        if self.data:

            if len(self.data):
                self.handlerDataBefore()
                self.handlerDataAfter()

            if len(self.dataOutput):
                self.handlerDataFinish()
                self.successSignal()
            else:
                errors = ['ERR_SEARCH_PARSE_0']
                self.errorsSignal(errors)

    def handlerDataBefore(self):

        for indexPage, page in enumerate(self.data):
            self.preparePageTableHead(indexPage, page)

        # for indexPage, page in enumerate(self.data):

        #    self.preparePageTableBody(indexPage, page)

    def handlerDataAfter(self):

        for indexPage, page in enumerate(self.data):
            if len(page):
                for indexRow, row in enumerate(page):
                    if not len(self.dataColumnsSearch):
                        result = self.searchHead(row)

                        if result is not None and "output" in result and "length" in result:
                            self.searchHeadAfter(result)

                    else:

                        result = self.searchHead(row, 2)

                        if result is not None and "output" in result and "length" in result:
                            self.searchHeadAfter(result)
                        else:

                            result = self.searchData(row)

                            if result is not None:
                                self.dataOutput.append(result)
                            else:
                                self.searchHeadReset()

                    del page[indexRow]

                    if len(page) == 0:
                        del self.data[indexPage]

                    self.handlerDataAfter()

                    break

    def handlerDataFinish(self):

        xml = minidom.Document()
        root = xml.createElement('root')
        xml.appendChild(root)
        elements = xml.createElement('elements')
        root.appendChild(elements)

        for row in self.dataOutput:

            type = 'element'
            empty = 0

            for cell in row:
                value = str(row[cell])

                if not len(value):
                    empty += 1

            if empty != len(row):
                if len(row) - empty == 1 and "name" in row:
                    if len(row["name"]):
                        type = 'category'

                element = xml.createElement('element')
                element.setAttribute('type', type)

                for column in self.dataColumnsResult:
                    if type == 'category' and column != 'name':
                        break

                    elementCell = xml.createElement(column)

                    if column in row:
                        elementCell.appendChild(xml.createTextNode(str(row[column])))

                    element.appendChild(elementCell)

                elements.appendChild(element)

        with open(self.file["out"], "wb") as file:
            file.write(xml.toprettyxml(indent="\t", encoding="utf-8"))
