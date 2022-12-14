import io
import json
import docx
import pandas
import pathlib
import subprocess

from module.base import BaseParser


class WordPaser(BaseParser):
    def word_handler(self, extention):
        if extention == '.doc':
            subprocess.call(
                ['soffice', '--headless', '--convert-to', 'docx', self.file["in"], '--outdir', self.directory["in"]])
            # old_path = self.file["in"]
            path = pathlib.Path(self.file["in"])
            self.file["in"] = path.with_name(path.stem + ".docx")

            # Os.unlink(old_path)

        doc = docx.Document(self.file["in"])
        tables = []

        for section in doc.sections:
            if not section.header:
                continue

            if not section.header.tables:
                continue

            for header in section.header.tables:
                tables.append(header)

        if len(doc.tables):
            for table in doc.tables:
                tables.append(table)

        if not len(tables):
            return

        self.word_load_tables(tables)

    def word_load_tables(self, tables):

        data = {}
        irow = 0

        # print('WordLoadTables')

        if len(tables):
            for table in tables:
                if not len(table.rows):
                    continue

                for row in table.rows:
                    if not len(row.cells):
                        continue

                    vcell = {}
                    icell = 0

                    for cell in row.cells:
                        vcell[icell] = cell.text
                        icell += 1

                    if not len(vcell):
                        continue

                    if irow == 0:
                        data[irow] = ";".join('"' + str(x) + '"' for x in range(0, len(vcell)))
                        irow += 1

                    data[irow] = ";".join('"' + str(x) + '"' for x in vcell.values())
                    irow += 1

        if not len(data):
            self.handler_data()
            return

        data = "\r\n".join(str(x) for x in data.values())
        data = pandas.read_csv(io.StringIO(data), sep=';')
        data = self.ps_cleaner(data)
        data = json.loads(data.to_json(orient="values"))

        if not data:
            self.handler_data()
            return

        data = self.prepare_data(data)

        if data:
            self.data.append(data)

        self.handler_data()
